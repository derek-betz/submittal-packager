"""Reporting utilities for Submittal Packager."""

from __future__ import annotations

from datetime import datetime
from pathlib import Path
from typing import Any, Dict, Iterable, List

from docx import Document
from jinja2 import Environment, FileSystemLoader

from .config import Config
from .models import ManifestEntry, ValidationMessage, ValidationResult


def _environment(template_path: Path) -> Environment:
    loader = FileSystemLoader(str(template_path.parent))
    return Environment(loader=loader, autoescape=False)


def _ensure_package_overview(
    manifest: List[ManifestEntry], package_overview: Dict[str, Any] | None
) -> Dict[str, Any]:
    if package_overview is not None:
        return package_overview

    totals_files = len(manifest)
    totals_pages = sum(entry.pages for entry in manifest)
    discipline_summary: Dict[str, Dict[str, int]] = {}
    for entry in manifest:
        discipline = entry.discipline or "UNASSIGNED"
        stats = discipline_summary.setdefault(discipline, {"files": 0, "pages": 0})
        stats["files"] += 1
        stats["pages"] += entry.pages

    return {
        "root": "",
        "totals": {"files": totals_files, "pages": totals_pages},
        "folders": [],
        "folder_summary": {},
        "discipline_summary": discipline_summary,
        "extension_summary": {},
        "generated": [],
    }


def generate_transmittal_docx(
    *,
    config: Config,
    manifest: List[ManifestEntry],
    stage: str,
    output_path: Path,
    generated_at: str,
    messages: ValidationResult,
    template_path: Path | None = None,
    package_overview: Dict[str, Any] | None = None,
) -> None:
    """Render a DOCX transmittal using python-docx."""

    if template_path is None:
        template_path = Path(config.templates.transmittal_docx)
    env = _environment(template_path)
    template = env.get_template(template_path.name)
    totals_files = len(manifest)
    totals_pages = sum(entry.pages for entry in manifest)
    overview = _ensure_package_overview(manifest, package_overview)
    rendered = template.render(
        project=config.project.dict(),
        stage=stage,
        generated_at=generated_at,
        totals={"files": totals_files, "pages": totals_pages},
        files=manifest,
        exceptions={
            "errors": [msg.text for msg in messages.errors],
            "warnings": [msg.text for msg in messages.warnings],
        },
        package_overview=overview,
    )

    document = Document()
    for line in rendered.splitlines():
        if line.strip() == "":
            document.add_paragraph("")
        else:
            document.add_paragraph(line)
    document.save(output_path)


def generate_html_report(
    *,
    config: Config,
    manifest: List[ManifestEntry],
    stage: str,
    output_path: Path,
    generated_at: str,
    messages: ValidationResult,
    template_path: Path | None = None,
    package_overview: Dict[str, Any] | None = None,
) -> None:
    """Render HTML validation report."""

    if template_path is None:
        template_path = Path(config.templates.report_html)
    env = _environment(template_path)
    template = env.get_template(template_path.name)

    overview = _ensure_package_overview(manifest, package_overview)

    html = template.render(
        project=config.project.dict(),
        stage=stage,
        generated_at=generated_at,
        totals={"files": len(manifest), "pages": sum(entry.pages for entry in manifest)},
        errors=[msg.text for msg in messages.errors],
        warnings=[msg.text for msg in messages.warnings],
        files=manifest,
        checksum_algo=config.packaging.checksum_algo,
        package_overview=overview,
    )
    output_path.write_text(html)


__all__ = ["generate_transmittal_docx", "generate_html_report"]
